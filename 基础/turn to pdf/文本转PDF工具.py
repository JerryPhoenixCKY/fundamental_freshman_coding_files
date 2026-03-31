import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import os
from pathlib import Path
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import re

# 尝试注册中文字体
def register_fonts():
    """注册系统字体以支持中文"""
    try:
        # Windows 常见字体路径
        font_paths = [
            "C:\\Windows\\Fonts\\simhei.ttf",  # 黑体
            "C:\\Windows\\Fonts\\simsun.ttf",  # 宋体
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                font_name = Path(font_path).stem
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    return font_name
                except:
                    continue
    except:
        pass
    return "Helvetica"


def clean_math_formula(text):
    """
    清理文本中的数学公式乱码
    将LaTeX格式的公式转换为可读的格式
    """
    # 替换常见的LaTeX公式符号
    replacements = {
        r'\alpha': 'α',
        r'\beta': 'β',
        r'\gamma': 'γ',
        r'\delta': 'δ',
        r'\epsilon': 'ε',
        r'\theta': 'θ',
        r'\lambda': 'λ',
        r'\mu': 'μ',
        r'\pi': 'π',
        r'\sigma': 'σ',
        r'\tau': 'τ',
        r'\phi': 'φ',
        r'\psi': 'ψ',
        r'\omega': 'ω',
        r'\times': '×',
        r'\div': '÷',
        r'\pm': '±',
        r'\infty': '∞',
        r'\approx': '≈',
        r'\neq': '≠',
        r'\leq': '≤',
        r'\geq': '≥',
        r'\sqrt': '√',
        r'\sum': '∑',
        r'\int': '∫',
    }
    
    for latex, symbol in replacements.items():
        text = text.replace(latex, symbol)
    
    return text


def text_to_pdf(text, output_path):
    """
    将文本转换为PDF文件
    """
    try:
        # 注册字体
        font_name = register_fonts()
        
        # 清理数学公式
        cleaned_text = clean_math_formula(text)
        
        # 创建PDF文档
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # 创建文本样式
        styles = getSampleStyleSheet()
        
        # 如果成功注册中文字体，使用中文字体
        if font_name != "Helvetica":
            style = ParagraphStyle(
                'CustomStyle',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                leading=16,
                encoding='utf-8'
            )
        else:
            style = styles['Normal']
        
        # 处理换行和段落
        story = []
        paragraphs = cleaned_text.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                # 处理段落中的多个换行
                lines = para.strip().split('\n')
                for line in lines:
                    if line.strip():
                        try:
                            p = Paragraph(line.strip(), style)
                            story.append(p)
                        except:
                            # 如果出现问题，使用纯文本
                            story.append(Paragraph(line.strip(), styles['Normal']))
                
                story.append(Spacer(1, 0.2*inch))
        
        # 生成PDF
        doc.build(story)
        return True
    
    except Exception as e:
        raise Exception(f"生成PDF时出错: {str(e)}")


def doc_to_pdf(doc_path, output_path):
    """
    将Word文档转换为PDF
    """
    try:
        from docx import Document
        
        # 读取Word文档
        doc = Document(doc_path)
        
        # 提取所有文本
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                all_text.append(' | '.join(row_text))
        
        full_text = '\n\n'.join(all_text)
        
        # 使用text_to_pdf生成PDF
        text_to_pdf(full_text, output_path)
        return True
    
    except ImportError:
        raise Exception("需要安装python-docx库。请运行: pip install python-docx")
    except Exception as e:
        raise Exception(f"转换文档时出错: {str(e)}")


class MainWindow:
    def __init__(self, root):
        self.root = root
        self.root.title("文本转PDF工具")
        self.root.geometry("400x300")
        self.root.resizable(False, False)
        
        # 标题
        title_label = tk.Label(
            root,
            text="文本转PDF工具",
            font=("微软雅黑", 18, "bold")
        )
        title_label.pack(pady=20)
        
        # 说明文本
        info_label = tk.Label(
            root,
            text="选择一个选项开始：",
            font=("微软雅黑", 12)
        )
        info_label.pack(pady=10)
        
        # 选项1按钮
        btn1 = tk.Button(
            root,
            text="选项1：输入文本生成PDF",
            font=("微软雅黑", 12),
            width=30,
            height=3,
            bg="#4CAF50",
            fg="white",
            command=self.option1
        )
        btn1.pack(pady=10)
        
        # 选项2按钮
        btn2 = tk.Button(
            root,
            text="选项2：选择doc文件转为PDF",
            font=("微软雅黑", 12),
            width=30,
            height=3,
            bg="#2196F3",
            fg="white",
            command=self.option2
        )
        btn2.pack(pady=10)
        
        # 退出按钮
        exit_btn = tk.Button(
            root,
            text="退出",
            font=("微软雅黑", 11),
            width=20,
            command=root.quit
        )
        exit_btn.pack(pady=20)
    
    def option1(self):
        """选项1：直接输入文本"""
        window = tk.Toplevel(self.root)
        window.title("输入文本")
        window.geometry("700x600")
        
        # 标题
        title = tk.Label(
            window,
            text="输入您复制的文本（包括数学公式）",
            font=("微软雅黑", 12, "bold")
        )
        title.pack(pady=10)
        
        # 文本输入框
        text_box = scrolledtext.ScrolledText(
            window,
            width=80,
            height=20,
            font=("微软雅黑", 10),
            wrap=tk.WORD
        )
        text_box.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # 按钮框架
        button_frame = tk.Frame(window)
        button_frame.pack(pady=10)
        
        def generate():
            text = text_box.get("1.0", tk.END)
            if not text.strip():
                messagebox.showwarning("警告", "请输入一些文本！")
                return
            
            # 文件保存对话框
            save_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF文件", "*.pdf"), ("全部文件", "*.*")],
                initialfile="输出文档.pdf"
            )
            
            if save_path:
                try:
                    text_to_pdf(text, save_path)
                    messagebox.showinfo("成功", f"PDF已生成：\n{save_path}")
                    window.destroy()
                except Exception as e:
                    messagebox.showerror("错误", str(e))
        
        # 生成PDF按钮
        generate_btn = tk.Button(
            button_frame,
            text="生成PDF",
            font=("微软雅黑", 11),
            width=15,
            bg="#4CAF50",
            fg="white",
            command=generate
        )
        generate_btn.pack(side=tk.LEFT, padx=5)
        
        # 取消按钮
        cancel_btn = tk.Button(
            button_frame,
            text="取消",
            font=("微软雅黑", 11),
            width=15,
            command=window.destroy
        )
        cancel_btn.pack(side=tk.LEFT, padx=5)
    
    def option2(self):
        """选项2：选择doc文件"""
        file_path = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[
                ("Word文档", "*.docx;*.doc"),
                ("DOCX文件", "*.docx"),
                ("DOC文件", "*.doc"),
                ("全部文件", "*.*")
            ]
        )
        
        if not file_path:
            return
        
        # 确定输出路径
        output_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF文件", "*.pdf"), ("全部文件", "*.*")],
            initialfile=Path(file_path).stem + ".pdf"
        )
        
        if output_path:
            try:
                messagebox.showinfo("进行中", "正在转换文档，请稍候...")
                doc_to_pdf(file_path, output_path)
                messagebox.showinfo("成功", f"PDF已生成：\n{output_path}")
            except Exception as e:
                messagebox.showerror("错误", str(e))


def main():
    root = tk.Tk()
    app = MainWindow(root)
    root.mainloop()


if __name__ == "__main__":
    main()
